#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
交通事故赔偿明细Word报告生成器
生成可直接给当事人看、可打印归档的赔偿计算报告
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 赔偿项目元数据（用于生成更专业的报告）
# ============================================================
COMPENSATION_META = {
    'medical_fee': {
        'name': '医疗费',
        'legal_basis': '《民法典》第1179条',
        'description': '按实际票据金额计算',
    },
    'hospital_meal_fee': {
        'name': '住院伙食补助费',
        'legal_basis': '《民法典》第1179条',
        'description': '参照国家机关一般工作人员出差伙食标准',
    },
    'nutrition_fee': {
        'name': '营养费',
        'legal_basis': '《民法典》第1179条',
        'description': '参照受害人伤残情况和医疗意见确定',
    },
    'nursing_fee': {
        'name': '护理费',
        'legal_basis': '《民法典》第1179条',
        'description': '参照当地护工从事同等级别护理的劳务报酬标准',
    },
    'lost_wage': {
        'name': '误工费',
        'legal_basis': '《民法典》第1179条',
        'description': '根据受害人误工时间和收入状况确定',
    },
    'traffic_fee': {
        'name': '交通费',
        'legal_basis': '《民法典》第1179条',
        'description': '按受害人及其必要的陪护人员就医实际发生的费用计算',
    },
    'disability_compensation': {
        'name': '残疾赔偿金',
        'legal_basis': '《民法典》第1179条/最高法人身损害赔偿解释',
        'description': '按城镇居民人均可支配收入×伤残系数×赔偿年限',
    },
    'death_compensation': {
        'name': '死亡赔偿金',
        'legal_basis': '《民法典》第1179条/最高法人身损害赔偿解释',
        'description': '按城镇居民人均可支配收入×20年（60岁以上递减）',
    },
    'funeral_fee': {
        'name': '丧葬费',
        'legal_basis': '《民法典》第1179条',
        'description': '按全省上年度职工月平均工资标准×6个月',
    },
    'dependent_living': {
        'name': '被扶养人生活费',
        'legal_basis': '《民法典》第1179条/最高法人身损害赔偿解释',
        'description': '计入残疾/死亡赔偿金，年赔偿总额不超过人均消费支出',
    },
    'mental_damage_fee': {
        'name': '精神损害抚慰金',
        'legal_basis': '《民法典》第1183条',
        'description': '因侵害人身权益造成严重精神损害的，可请求赔偿',
    },
    'property_damage': {
        'name': '财产损失',
        'legal_basis': '《民法典》第1184条',
        'description': '按照损失发生时的市场价格计算',
    },
    'other_fee': {
        'name': '其他费用',
        'legal_basis': '《民法典》第1179条',
        'description': '鉴定费、后续治疗费等',
    },
}


def generate_compensation_report_docx(
    case_data: dict,
    result: dict,
    output_path: str = None,
    plaintiff_name: str = None,
    defendant_names: list = None,
    accident_date: str = None,
) -> str:
    """
    生成赔偿明细Word报告
    
    参数:
        case_data: 案件基础数据（用于提取计算参数）
        result: calculate()返回的计算结果字典
        output_path: 输出文件路径
        plaintiff_name: 原告姓名
        defendant_names: 被告姓名列表
        accident_date: 事故日期
    
    返回: 生成的文件路径
    """
    if output_path is None:
        output_path = '赔偿明细报告.docx'

    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ========== 标题 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('交通事故赔偿明细计算报告')
    _set_font(run, '黑体', 18, bold=True)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.space_before = Pt(0)

    # ========== 案件基本信息 ==========
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(subtitle.add_run('── 案件基本信息 ──'), '宋体', 10, color=RGBColor(128, 128, 128))
    subtitle.paragraph_format.space_after = Pt(6)

    # 基本信息表格
    info_table = doc.add_table(rows=5, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'

    injury_text = {'death': '死亡', 'disability': '伤残', 'property': '仅财产损失'}.get(
        case_data.get('injury_type', ''), '未知')
    case_type_text = {'civil': '纯民事', 'criminal_attached': '刑事附带民事'}.get(
        case_data.get('case_type', ''), '未知')
    region_text = '城镇' if case_data.get('region', 'urban') == 'urban' else '农村'

    info_data = [
        ['原告', plaintiff_name or case_data.get('plaintiff_name', '×××'), '被告', '、'.join(defendant_names) if defendant_names else case_data.get('defendant_text', '×××')],
        ['赔偿标准', f'河北省{case_data.get("standard_year", "2024")}年标准（{region_text}）', '案件类型', case_type_text],
        ['伤情类型', injury_text, '责任比例', f'{case_data.get("liability_type", "全责")}（{_get_ratio(case_data)*100:.0f}%）'],
        ['住院天数', f'{case_data.get("hospital_days", 0)}天', '事故日期', accident_date or case_data.get('accident_date', '××××年××月××日')],
        ['伤残等级', f'{case_data.get("disability_grade", "—")}级' if case_data.get('injury_type') == 'disability' else '—', '受害人年龄', f'{case_data.get("plaintiff_age", "—")}岁'],
    ]

    for i, row_data in enumerate(info_data):
        for j, text in enumerate(row_data):
            cell = info_table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            if j % 2 == 0:  # 标签列
                _set_font(run, '宋体', 10, bold=True, color=RGBColor(51, 51, 51))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_shading(cell, 'F2F2F2')
            else:  # 内容列
                _set_font(run, '宋体', 10)

    doc.add_paragraph()  # 空行

    # ========== 赔偿项目明细 ==========
    section_title = doc.add_paragraph()
    _set_font(section_title.add_run('一、赔偿项目明细'), '黑体', 14, bold=True)
    section_title.paragraph_format.space_after = Pt(6)

    # 确定赔偿项目列表
    items = _build_items(case_data, result)

    # 主表格: 序号 | 赔偿项目 | 计算公式 | 金额(元) | 法律依据
    table = doc.add_table(rows=1 + len(items) + 1, cols=5)  # +1 for total
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 列宽
    col_widths = [Cm(1.2), Cm(3.5), Cm(5.5), Cm(3.0), Cm(3.4)]

    # 表头
    headers = ['序号', '赔偿项目', '计算公式', '金额（元）', '法律依据']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_font(run, '宋体', 10, bold=True, color=RGBColor(255, 255, 255))
        _set_cell_shading(cell, '4472C4')

    for i, cell in enumerate(table.rows[0].cells):
        cell.width = col_widths[i]

    # 填充数据
    for idx, item in enumerate(items):
        row = table.rows[idx + 1]
        data = [
            item['seq'],
            item['name'],
            item['formula'],
            f"{item['amount']:,.2f}" if item['amount'] != 0 else '—',
            item.get('legal', ''),
        ]
        for i, text in enumerate(data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            if i in (0, 3):  # 居中
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            _set_font(run, '宋体', 10)
            if item.get('highlight'):
                run.font.color.rgb = RGBColor(192, 0, 0)  # 红色标记不支持项目

        # 交替行底色
        if idx % 2 == 1:
            for cell in row.cells:
                _set_cell_shading(cell, 'F2F7FC')

    # 合计行
    total_row = table.rows[len(items) + 1]
    total_row.cells[0].text = ''
    total_row.cells[1].text = ''
    total_row.cells[2].text = ''
    # 合并前3列
    total_row.cells[0].merge(total_row.cells[2])
    p = total_row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('标的总额')
    _set_font(run, '宋体', 11, bold=True)

    p2 = total_row.cells[3].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{result['total_amount']:,.2f}")
    _set_font(run2, '宋体', 11, bold=True, color=RGBColor(192, 0, 0))

    total_row.cells[4].text = ''

    # 合计行底色
    for cell in total_row.cells:
        _set_cell_shading(cell, 'E2EFDA')

    doc.add_paragraph()  # 空行

    # ========== 保险赔付拆分 ==========
    section_title2 = doc.add_paragraph()
    _set_font(section_title2.add_run('二、保险赔付拆分'), '黑体', 14, bold=True)
    section_title2.paragraph_format.space_after = Pt(6)

    ins = result.get('insurance_split', {})
    liability_type = case_data.get('liability_type', '全责')
    ratio = _get_ratio(case_data)

    ins_table = doc.add_table(rows=5, cols=3)
    ins_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ins_table.style = 'Table Grid'

    ins_headers = ['赔付方', '计算说明', '金额（元）']
    for i, h in enumerate(ins_headers):
        cell = ins_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, '宋体', 10, bold=True, color=RGBColor(255, 255, 255))
        _set_cell_shading(cell, '4472C4')

    ins_data = [
        ['交强险', '限额内全额赔付（不分责任比例）', ins.get('compulsory', 0)],
        [f'商业三者险（{liability_type}{ratio*100:.0f}%）', f'超出交强险部分×{ratio*100:.0f}%', ins.get('commercial', 0)],
        [f'原告自行承担（{(1-ratio)*100:.0f}%）', f'超出交强险部分×{(1-ratio)*100:.0f}%', ins.get('self_bear', 0)],
        ['保险公司合计', '交强险 + 商业三者险', ins.get('total_insurance', 0)],
    ]

    for i, (payer, desc, amount) in enumerate(ins_data):
        row = ins_table.rows[i + 1]
        for j, text in enumerate([payer, desc, f'{amount:,.2f}']):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            if j == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            _set_font(run, '宋体', 10)
            if i == 3:  # 合计行
                run.font.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)
                _set_cell_shading(cell, 'E2EFDA')

    # 合计行底色
    for cell in ins_table.rows[4].cells:
        _set_cell_shading(cell, 'E2EFDA')

    doc.add_paragraph()

    # ========== 交强险分项拆解 ==========
    section_title3 = doc.add_paragraph()
    _set_font(section_title3.add_run('三、交强险分项限额'), '黑体', 14, bold=True)
    section_title3.paragraph_format.space_after = Pt(6)

    compulsory_table = doc.add_table(rows=4, cols=4)
    compulsory_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    compulsory_table.style = 'Table Grid'

    comp_headers = ['分项', '限额（元）', '本项目损失（元）', '可赔付（元）']
    for i, h in enumerate(comp_headers):
        cell = compulsory_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, '宋体', 10, bold=True, color=RGBColor(255, 255, 255))
        _set_cell_shading(cell, '4472C4')

    # 计算各分项
    medical_items = result.get('medical_fee', 0) + result.get('hospital_meal_fee', 0) + result.get('nutrition_fee', 0)
    death_disability_items = result.get('nursing_fee', 0) + result.get('lost_wage', 0) + result.get('traffic_fee', 0) + \
                             result.get('disability_compensation', 0) + result.get('death_compensation', 0) + \
                             result.get('funeral_fee', 0) + result.get('dependent_living', 0) + \
                             result.get('mental_damage_fee', 0)
    property_items = result.get('property_damage', 0)

    comp_data = [
        ['医疗费用分项', '18,000', f'{medical_items:,.2f}', f'{min(medical_items, 18000):,.2f}'],
        ['死亡伤残分项', '180,000', f'{death_disability_items:,.2f}', f'{min(death_disability_items, 180000):,.2f}'],
        ['财产损失分项', '2,000', f'{property_items:,.2f}', f'{min(property_items, 2000):,.2f}'],
    ]

    for i, row_data in enumerate(comp_data):
        for j, text in enumerate(row_data):
            cell = compulsory_table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            _set_font(run, '宋体', 10)

    doc.add_paragraph()

    # ========== 重要提示 ==========
    section_title4 = doc.add_paragraph()
    _set_font(section_title4.add_run('四、重要提示'), '黑体', 14, bold=True)
    section_title4.paragraph_format.space_after = Pt(6)

    tips = _build_tips(case_data, result)
    for tip in tips:
        p = doc.add_paragraph()
        run = p.add_run(f'• {tip}')
        _set_font(run, '宋体', 10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    # ========== 落款 ==========
    doc.add_paragraph()
    doc.add_paragraph()

    note_p = doc.add_paragraph()
    _set_font(note_p.add_run('本报告由交通事故赔偿计算系统自动生成，仅供参考。具体赔偿金额以法院判决为准。'), '宋体', 9, color=RGBColor(128, 128, 128))

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(date_p.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}'), '宋体', 10)

    # 保存
    doc.save(output_path)
    return output_path


def _build_items(case_data: dict, result: dict) -> list:
    """构建赔偿项目列表"""
    items = []
    seq = 1

    # 1. 医疗费
    if result.get('medical_fee', 0) > 0:
        items.append({
            'seq': str(seq), 'name': '医疗费',
            'formula': '按实际票据',
            'amount': result['medical_fee'],
            'legal': COMPENSATION_META['medical_fee']['legal_basis'],
        })
        seq += 1

    # 2. 住院伙食补助费
    if result.get('hospital_meal_fee', 0) > 0:
        days = case_data.get('hospital_days', 0)
        daily = case_data.get('hospital_meal_daily') or 100
        items.append({
            'seq': str(seq), 'name': '住院伙食补助费',
            'formula': f'{days}天 × {daily}元/天',
            'amount': result['hospital_meal_fee'],
            'legal': COMPENSATION_META['hospital_meal_fee']['legal_basis'],
        })
        seq += 1

    # 3. 营养费
    if result.get('nutrition_fee', 0) > 0:
        days = case_data.get('nutrition_days') or case_data.get('hospital_days', 0)
        daily = case_data.get('nutrition_daily') or 30
        items.append({
            'seq': str(seq), 'name': '营养费',
            'formula': f'{days}天 × {daily}元/天',
            'amount': result['nutrition_fee'],
            'legal': COMPENSATION_META['nutrition_fee']['legal_basis'],
        })
        seq += 1

    # 4. 护理费
    if result.get('nursing_fee', 0) > 0:
        days = case_data.get('nursing_days') or case_data.get('hospital_days', 0)
        persons = case_data.get('nursing_persons', 1)
        daily = case_data.get('nursing_custom_daily') or 130
        formula = f'{days}天 × {persons}人 × {daily}元/天' if persons > 1 else f'{days}天 × {daily}元/天'
        items.append({
            'seq': str(seq), 'name': '护理费',
            'formula': formula,
            'amount': result['nursing_fee'],
            'legal': COMPENSATION_META['nursing_fee']['legal_basis'],
        })
        seq += 1

    # 5. 误工费
    if result.get('lost_wage', 0) > 0:
        days = case_data.get('lost_work_days') or case_data.get('hospital_days', 0)
        daily = case_data.get('lost_work_daily') or 178
        items.append({
            'seq': str(seq), 'name': '误工费',
            'formula': f'{days}天 × {daily}元/天',
            'amount': result['lost_wage'],
            'legal': COMPENSATION_META['lost_wage']['legal_basis'],
        })
        seq += 1

    # 6. 交通费
    if result.get('traffic_fee', 0) > 0:
        items.append({
            'seq': str(seq), 'name': '交通费',
            'formula': '按实际票据',
            'amount': result['traffic_fee'],
            'legal': COMPENSATION_META['traffic_fee']['legal_basis'],
        })
        seq += 1

    # 7. 残疾/死亡赔偿金
    injury_type = case_data.get('injury_type', '')
    if injury_type == 'death':
        age = case_data.get('plaintiff_age', 0)
        standard_year = case_data.get('standard_year', '2024')
        income = _get_income(case_data)
        years = 20 if age < 60 else max(5, 20 - (age - 60))
        items.append({
            'seq': str(seq), 'name': '死亡赔偿金',
            'formula': f'{income:,.0f}元/年 × {years}年',
            'amount': result.get('death_compensation', 0),
            'legal': COMPENSATION_META['death_compensation']['legal_basis'],
        })
        seq += 1

        items.append({
            'seq': f'{seq-1}-1', 'name': '丧葬费',
            'formula': '全省上年度职工月平均工资 × 6个月',
            'amount': result.get('funeral_fee', 0),
            'legal': COMPENSATION_META['funeral_fee']['legal_basis'],
        })

    elif injury_type == 'disability':
        grade = case_data.get('disability_grade', 10)
        age = case_data.get('plaintiff_age', 0)
        income = _get_income(case_data)
        ratio_map = {10: 0.10, 9: 0.20, 8: 0.30, 7: 0.40, 6: 0.50,
                     5: 0.60, 4: 0.70, 3: 0.80, 2: 0.90, 1: 1.0}
        ratio = ratio_map.get(grade, 0.10)
        years = 20 if age < 60 else max(5, 20 - (age - 60))
        items.append({
            'seq': str(seq), 'name': '残疾赔偿金',
            'formula': f'{income:,.0f}元/年 × {years}年 × {ratio*100:.0f}%',
            'amount': result.get('disability_compensation', 0),
            'legal': COMPENSATION_META['disability_compensation']['legal_basis'],
        })
        seq += 1

    # 7-2. 被扶养人生活费
    if result.get('dependent_living', 0) > 0:
        items.append({
            'seq': f'{seq-1}-2', 'name': '被扶养人生活费',
            'formula': '分段计算（年赔偿总额≤人均消费支出）',
            'amount': result['dependent_living'],
            'legal': COMPENSATION_META['dependent_living']['legal_basis'],
        })

    # 8. 精神损害抚慰金
    mental = result.get('mental_damage_fee', 0)
    if mental > 0:
        items.append({
            'seq': str(seq), 'name': '精神损害抚慰金',
            'formula': f'{"死亡" if injury_type == "death" else f"{case_data.get("disability_grade", "")}级伤残"}标准',
            'amount': mental,
            'legal': COMPENSATION_META['mental_damage_fee']['legal_basis'],
        })
        seq += 1
    elif case_data.get('case_type') == 'criminal_attached':
        items.append({
            'seq': str(seq), 'name': '精神损害抚慰金',
            'formula': '❌ 刑事附带民事不支持（刑诉法解释第192条）',
            'amount': 0,
            'legal': COMPENSATION_META['mental_damage_fee']['legal_basis'],
            'highlight': True,
        })
        seq += 1

    # 9. 财产损失
    if result.get('property_damage', 0) > 0:
        items.append({
            'seq': str(seq), 'name': '财产损失',
            'formula': '车辆定损/维修费用',
            'amount': result['property_damage'],
            'legal': COMPENSATION_META['property_damage']['legal_basis'],
        })
        seq += 1

    # 10. 其他费用
    if result.get('other_fee', 0) > 0:
        items.append({
            'seq': str(seq), 'name': '其他费用',
            'formula': case_data.get('other_fee_desc', '鉴定费、后续治疗费等'),
            'amount': result['other_fee'],
            'legal': COMPENSATION_META['other_fee']['legal_basis'],
        })

    return items


def _build_tips(case_data: dict, result: dict) -> list:
    """构建重要提示"""
    tips = []
    injury_type = case_data.get('injury_type', '')
    case_type = case_data.get('case_type', '')
    liability_type = case_data.get('liability_type', '全责')

    if case_type == 'criminal_attached':
        tips.append('刑事附带民事案件不支持精神损害抚慰金（刑诉法解释第192条），如需主张需另行提起民事诉讼')

    if injury_type == 'death':
        tips.append('死亡赔偿金按20年计算，60周岁以上每增加1岁减少1年，最低5年')

    if injury_type == 'disability':
        grade = case_data.get('disability_grade', 10)
        tips.append(f'伤残等级{grade}级，赔偿系数{int({10:10,9:20,8:30,7:40,6:50,5:60,4:70,3:80,2:90,1:100}.get(grade, 10))}%，以司法鉴定意见为准')

    if liability_type != '全责':
        ratio = _get_ratio(case_data)
        tips.append(f'被告负{liability_type}，商业险部分按{ratio*100:.0f}%比例赔付，原告自行承担{(1-ratio)*100:.0f}%')

    tips.append('交强险在限额内全额赔付，不按责任比例划分')
    tips.append('赔偿标准以一审法庭辩论终结时上一年度统计数据为准')

    if result.get('dependent_living', 0) > 0:
        tips.append('被扶养人生活费计入残疾/死亡赔偿金，年赔偿总额不超过人均消费支出')

    tips.append('非医保用药部分可能由侵权人承担，不在保险赔付范围内')

    return tips


def _get_ratio(case_data: dict) -> float:
    """获取责任比例"""
    ratio_map = {'全责': 1.0, '主责': 0.7, '同责': 0.5, '次责': 0.3, '无责': 0.0}
    return ratio_map.get(case_data.get('liability_type', '全责'), 1.0)


def _get_income(case_data: dict) -> float:
    """获取收入标准"""
    from lawsuit_generator_v2 import HEBEI_STANDARD
    std_year = case_data.get('standard_year', '2024')
    std = HEBEI_STANDARD.get(std_year, HEBEI_STANDARD['2024'])
    region = case_data.get('region', 'urban')
    return std['urban_income'] if region == 'urban' else std['rural_income']


def _set_font(run, font_name, size, bold=False, color=RGBColor(0, 0, 0)):
    """统一设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def _set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading)


# ============================================================
# 测试
# ============================================================
if __name__ == '__main__':
    import sys
    sys.path.insert(0, '交通事故案件工作流')
    from lawsuit_generator_v2 import CaseData, HEBEI_STANDARD

    # === 冉某江案 ===
    print("生成冉某江案赔偿报告...")

    case1 = CaseData()
    case1.case_type = CaseData.CASE_TYPE_CRIMINAL_ATTACHED
    case1.injury_type = CaseData.INJURY_DEATH
    case1.standard_year = '2023'
    case1.liability_type = '主责'
    case1.hospital_days = 11
    case1.medical_fee = 102309.83
    case1.nursing_days = 11
    case1.nursing_custom_daily = 139
    case1.nutrition_days = 11
    case1.nutrition_daily = 20
    case1.hospital_meal_daily = 50
    case1.lost_work_days = 11
    case1.lost_work_daily = 178
    case1.traffic_fee = 3000
    case1.property_damage = 2000
    case1.plaintiff_age = 55

    result1 = case1.calculate()

    case1_dict = {
        'injury_type': case1.injury_type,
        'case_type': case1.case_type,
        'standard_year': case1.standard_year,
        'region': case1.region,
        'liability_type': case1.liability_type,
        'hospital_days': case1.hospital_days,
        'hospital_meal_daily': case1.hospital_meal_daily,
        'nutrition_days': case1.nutrition_days,
        'nutrition_daily': case1.nutrition_daily,
        'nursing_days': case1.nursing_days,
        'nursing_persons': case1.nursing_persons,
        'nursing_custom_daily': case1.nursing_custom_daily,
        'lost_work_days': case1.lost_work_days,
        'lost_work_daily': case1.lost_work_daily,
        'disability_grade': case1.disability_grade,
        'plaintiff_age': case1.plaintiff_age,
        'dependents': case1.dependents,
        'other_fee_desc': case1.other_fee_desc,
    }

    output = generate_compensation_report_docx(
        case1_dict, result1,
        output_path='交通事故案件工作流/赔偿明细报告_冉某江案.docx',
        plaintiff_name='冉某江',
        defendant_names=['潘某坤', '中国某某财产保险股份有限公司邢台市中心支公司'],
        accident_date='2023年10月28日',
    )
    print(f"✓ 赔偿报告已生成: {output}")
    print(f"  标的总额: {result1['total_amount']:,.2f}元")
    print(f"  保险公司合计: {result1['insurance_split']['total_insurance']:,.2f}元")

    # === 李明伤残案 ===
    print("\n生成李明伤残案赔偿报告...")

    case2 = CaseData()
    case2.case_type = CaseData.CASE_TYPE_CIVIL
    case2.injury_type = CaseData.INJURY_DISABILITY
    case2.standard_year = '2024'
    case2.liability_type = '全责'
    case2.hospital_days = 45
    case2.medical_fee = 52340.80
    case2.nursing_days = 45
    case2.nursing_persons = 1
    case2.nutrition_days = 60
    case2.lost_work_days = 180
    case2.lost_work_daily = 178
    case2.traffic_fee = 680
    case2.disability_grade = 10
    case2.plaintiff_age = 36

    result2 = case2.calculate()

    case2_dict = {
        'injury_type': case2.injury_type,
        'case_type': case2.case_type,
        'standard_year': case2.standard_year,
        'region': case2.region,
        'liability_type': case2.liability_type,
        'hospital_days': case2.hospital_days,
        'hospital_meal_daily': case2.hospital_meal_daily,
        'nutrition_days': case2.nutrition_days,
        'nutrition_daily': case2.nutrition_daily,
        'nursing_days': case2.nursing_days,
        'nursing_persons': case2.nursing_persons,
        'nursing_custom_daily': case2.nursing_custom_daily,
        'lost_work_days': case2.lost_work_days,
        'lost_work_daily': case2.lost_work_daily,
        'disability_grade': case2.disability_grade,
        'plaintiff_age': case2.plaintiff_age,
        'dependents': case2.dependents,
        'other_fee_desc': case2.other_fee_desc,
    }

    output2 = generate_compensation_report_docx(
        case2_dict, result2,
        output_path='交通事故案件工作流/赔偿明细报告_李明伤残案.docx',
        plaintiff_name='李明',
        defendant_names=['张某某', '某某保险股份有限公司'],
        accident_date='2024年6月15日',
    )
    print(f"✓ 赔偿报告已生成: {output2}")
    print(f"  标的总额: {result2['total_amount']:,.2f}元")
