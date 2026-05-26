#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
交通事故案件证据清单自动生成器
根据案件类型、伤情、费用项目自动推导所需证据，生成标准Word格式证据清单
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ============================================================
# 证据规则库：什么案件类型需要什么证据
# ============================================================
EVIDENCE_RULES = {
    # === 通用证据（所有案件都需要） ===
    'common': [
        {
            'name': '原告身份证复印件',
            'source': '原告',
            'purpose': '证明原告主体资格',
            'copies': 1,
            'condition': lambda c: True,
        },
        {
            'name': '道路交通事故认定书',
            'source': '公安交管部门',
            'purpose': '证明事故发生经过及责任划分',
            'copies': 1,
            'condition': lambda c: True,
        },
        {
            'name': '被告驾驶证、行驶证复印件',
            'source': '公安交管部门',
            'purpose': '证明被告身份及驾驶资格',
            'copies': 1,
            'condition': lambda c: True,
        },
        {
            'name': '机动车交强险保单',
            'source': '保险公司',
            'purpose': '证明肇事车辆投保交强险的事实',
            'copies': 1,
            'condition': lambda c: True,
        },
        {
            'name': '机动车商业三者险保单',
            'source': '保险公司',
            'purpose': '证明肇事车辆投保商业三者险的事实',
            'copies': 1,
            'condition': lambda c: True,
        },
    ],

    # === 医疗类证据 ===
    'medical': [
        {
            'name': '门诊病历',
            'source': '就诊医院',
            'purpose': '证明伤者门诊就诊及伤情',
            'copies': 1,
            'condition': lambda c: c.get('medical_fee', 0) > 0,
        },
        {
            'name': '住院病历（含入院记录、出院记录、手术记录、长期医嘱、临时医嘱）',
            'source': '就诊医院',
            'purpose': '证明伤者住院治疗经过及伤情',
            'copies': 1,
            'condition': lambda c: c.get('hospital_days', 0) > 0,
        },
        {
            'name': '诊断证明书',
            'source': '就诊医院',
            'purpose': '证明伤者伤情诊断及治疗建议',
            'copies': 1,
            'condition': lambda c: c.get('medical_fee', 0) > 0,
        },
        {
            'name': '医疗费发票及费用清单',
            'source': '就诊医院',
            'purpose': f'证明医疗费支出',
            'copies': 1,
            'condition': lambda c: c.get('medical_fee', 0) > 0,
            'amount_ref': 'medical_fee',
        },
        {
            'name': '住院费用明细清单',
            'source': '就诊医院',
            'purpose': '证明住院期间用药及治疗费用的合理性',
            'copies': 1,
            'condition': lambda c: c.get('hospital_days', 0) > 0,
        },
    ],

    # === 护理费证据 ===
    'nursing': [
        {
            'name': '护理人身份证复印件及收入证明',
            'source': '护理人工作单位',
            'purpose': '证明护理人身份及护理费计算依据',
            'copies': 1,
            'condition': lambda c: c.get('nursing_days', 0) > 0 and c.get('nursing_custom_daily') is not None,
        },
        {
            'name': '护理证明（医院出具需陪护证明）',
            'source': '就诊医院',
            'purpose': '证明住院期间需要护理的事实',
            'copies': 1,
            'condition': lambda c: c.get('nursing_days', 0) > 0,
        },
    ],

    # === 误工费证据 ===
    'lost_wage': [
        {
            'name': '误工证明（含停发工资证明）',
            'source': '原告工作单位',
            'purpose': '证明原告因伤误工及收入减少的事实',
            'copies': 1,
            'condition': lambda c: c.get('lost_work_days', 0) > 0,
        },
        {
            'name': '劳动合同、社保缴纳记录',
            'source': '原告工作单位/社保机构',
            'purpose': '证明原告劳动关系及收入情况',
            'copies': 1,
            'condition': lambda c: c.get('lost_work_days', 0) > 0 and c.get('lost_work_daily') is not None,
        },
        {
            'name': '事故前12个月工资银行流水',
            'source': '银行',
            'purpose': '证明原告实际收入水平，作为误工费计算依据',
            'copies': 1,
            'condition': lambda c: c.get('lost_work_days', 0) > 0 and c.get('lost_work_daily') is not None,
        },
    ],

    # === 交通费证据 ===
    'traffic': [
        {
            'name': '交通费票据',
            'source': '交通运输部门',
            'purpose': '证明原告因就医、转院等产生的交通费用',
            'copies': 1,
            'condition': lambda c: c.get('traffic_fee', 0) > 0,
            'amount_ref': 'traffic_fee',
        },
    ],

    # === 伤残鉴定证据 ===
    'disability': [
        {
            'name': '司法鉴定意见书',
            'source': '司法鉴定机构',
            'purpose': '证明伤残等级、三期（误工期、护理期、营养期）评定',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'disability',
        },
        {
            'name': '鉴定费发票',
            'source': '司法鉴定机构',
            'purpose': '证明鉴定费支出',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'disability',
        },
    ],

    # === 死亡案件证据 ===
    'death': [
        {
            'name': '居民死亡医学证明（推断）书',
            'source': '医院/公安机关',
            'purpose': '证明受害人死亡的事实及原因',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'death',
        },
        {
            'name': '死亡殡葬证',
            'source': '民政部门',
            'purpose': '证明受害人已办理死亡殡葬手续',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'death',
        },
        {
            'name': '户口本（全户）复印件',
            'source': '原告',
            'purpose': '证明原告与死者的亲属关系，作为赔偿权利人主体资格的依据',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'death',
        },
        {
            'name': '法定继承人关系证明',
            'source': '户籍所在地派出所/村委会',
            'purpose': '证明原告为死者第一顺序法定继承人',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'death',
        },
        {
            'name': '丧葬费票据',
            'source': '殡仪馆',
            'purpose': '证明丧葬费实际支出',
            'copies': 1,
            'condition': lambda c: c.get('injury_type') == 'death',
        },
    ],

    # === 财产损失证据 ===
    'property': [
        {
            'name': '车辆损失定损单',
            'source': '保险公司/评估机构',
            'purpose': '证明车辆损失金额',
            'copies': 1,
            'condition': lambda c: c.get('property_damage', 0) > 0,
        },
        {
            'name': '车辆维修发票及维修清单',
            'source': '维修厂',
            'purpose': '证明车辆维修费用',
            'copies': 1,
            'condition': lambda c: c.get('property_damage', 0) > 0,
        },
    ],

    # === 被扶养人证据 ===
    'dependent': [
        {
            'name': '被扶养人身份证复印件',
            'source': '原告',
            'purpose': '证明被扶养人身份',
            'copies': 1,
            'condition': lambda c: len(c.get('dependents', [])) > 0,
        },
        {
            'name': '被扶养人与受害人关系证明',
            'source': '户籍所在地派出所/村委会',
            'purpose': '证明被扶养人与受害人的亲属关系',
            'copies': 1,
            'condition': lambda c: len(c.get('dependents', [])) > 0,
        },
        {
            'name': '被扶养人无劳动能力证明/在校证明',
            'source': '社保机构/学校',
            'purpose': '证明被扶养人无劳动能力或尚在求学期间',
            'copies': 1,
            'condition': lambda c: len(c.get('dependents', [])) > 0,
        },
    ],

    # === 刑事附带民事额外证据 ===
    'criminal': [
        {
            'name': '刑事判决书',
            'source': '人民法院',
            'purpose': '证明被告已被追究刑事责任，本案系刑事附带民事',
            'copies': 1,
            'condition': lambda c: c.get('case_type') == 'criminal_attached',
        },
    ],
}


def collect_evidence(case_data: dict) -> list:
    """
    根据案件数据自动推导所需证据清单
    返回: [{name, source, purpose, copies, amount_ref}]
    """
    evidence = []
    seen = set()  # 去重

    # 按类别顺序收集
    category_order = ['common', 'medical', 'nursing', 'lost_wage', 'traffic',
                      'disability', 'death', 'property', 'dependent', 'criminal']

    for category in category_order:
        rules = EVIDENCE_RULES.get(category, [])
        for rule in rules:
            if rule['name'] in seen:
                continue
            if rule['condition'](case_data):
                evidence.append({
                    'name': rule['name'],
                    'source': rule['source'],
                    'purpose': rule['purpose'],
                    'copies': rule.get('copies', 1),
                    'amount_ref': rule.get('amount_ref'),
                })
                seen.add(rule['name'])

    # 追加用户自定义证据
    for custom in case_data.get('custom_evidence', []):
        if isinstance(custom, str):
            evidence.append({
                'name': custom,
                'source': '原告',
                'purpose': '（待补充）',
                'copies': 1,
                'amount_ref': None,
            })
        elif isinstance(custom, dict):
            evidence.append({
                'name': custom.get('name', ''),
                'source': custom.get('source', '原告'),
                'purpose': custom.get('purpose', '（待补充）'),
                'copies': custom.get('copies', 1),
                'amount_ref': custom.get('amount_ref'),
            })

    return evidence


def generate_evidence_list_docx(case_data: dict, output_path: str = None,
                                 plaintiff_name: str = None,
                                 defendant_names: list = None,
                                 court: str = None) -> str:
    """
    生成标准证据清单Word文档
    
    参数:
        case_data: 案件数据字典（含 injury_type, medical_fee, hospital_days 等）
        output_path: 输出文件路径
        plaintiff_name: 原告姓名
        defendant_names: 被告姓名列表
        court: 管辖法院
    
    返回: 生成的文件路径
    """
    evidence = collect_evidence(case_data)

    if output_path is None:
        output_path = '证据清单.docx'

    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ========== 标题 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('证  据  清  单')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    # 段后间距
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.space_before = Pt(0)

    # ========== 案件信息 ==========
    info_para = doc.add_paragraph()
    info_para.paragraph_format.space_after = Pt(6)
    info_para.paragraph_format.space_before = Pt(0)

    plaintiff_text = plaintiff_name or case_data.get('plaintiff_name', '×××')
    defendant_text = '、'.join(defendant_names) if defendant_names else case_data.get('defendant_text', '×××')
    court_text = court or case_data.get('court', '×××人民法院')

    run = info_para.add_run(f'原告：{plaintiff_text}')
    _set_font(run, '宋体', Pt(12))

    info_para2 = doc.add_paragraph()
    info_para2.paragraph_format.space_after = Pt(6)
    info_para2.paragraph_format.space_before = Pt(0)
    run = info_para2.add_run(f'被告：{defendant_text}')
    _set_font(run, '宋体', Pt(12))

    info_para3 = doc.add_paragraph()
    info_para3.paragraph_format.space_after = Pt(12)
    info_para3.paragraph_format.space_before = Pt(0)
    run = info_para3.add_run(f'案由：机动车交通事故责任纠纷')
    _set_font(run, '宋体', Pt(12))

    # ========== 证据表格 ==========
    # 列: 序号 | 证据名称 | 份数 | 页数 | 证据形式 | 证明目的
    table = doc.add_table(rows=1 + len(evidence), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    headers = ['序号', '证据名称', '份数', '页数', '证据形式', '证明目的']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_font(run, '宋体', Pt(10), bold=True)
        # 表头背景
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D9E2F3')
        cell._element.get_or_add_tcPr().append(shading)

    # 列宽设置
    col_widths = [Cm(1.2), Cm(6.5), Cm(1.2), Cm(1.2), Cm(2.0), Cm(4.5)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # 填充数据
    for idx, ev in enumerate(evidence):
        row = table.rows[idx + 1]
        data = [
            str(idx + 1),
            ev['name'],
            str(ev.get('copies', 1)),
            '',  # 页数留空，用户手填
            '书证',  # 默认书证
            ev['purpose'],
        ]
        for i, text in enumerate(data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            if i in (0, 2, 3, 4):  # 居中列
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            _set_font(run, '宋体', Pt(10))

    # ========== 落款 ==========
    doc.add_paragraph()  # 空行

    note_para = doc.add_paragraph()
    run = note_para.add_run('注：以上证据均为复印件，原件待庭审时出示。')
    _set_font(run, '宋体', Pt(10))
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    # 提交人
    submit_para = doc.add_paragraph()
    submit_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = submit_para.add_run(f'提交人：{plaintiff_text}')
    _set_font(run, '宋体', Pt(12))

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run(f'{datetime.now().strftime("%Y年%m月%d日")}')
    _set_font(run, '宋体', Pt(12))

    # 保存
    doc.save(output_path)
    return output_path


def generate_evidence_list_text(case_data: dict) -> str:
    """
    生成要素式起诉状中段落231所需的简版证据清单文本
    格式: "1. xxx；2. xxx；3. xxx。"
    """
    evidence = collect_evidence(case_data)
    items = [ev['name'] for ev in evidence]
    return '；'.join([f'{i+1}. {item}' for i, item in enumerate(items)]) + '。'


def _set_font(run, font_name, size, bold=False):
    """统一设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size if isinstance(size, Pt) else Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


# ============================================================
# 测试
# ============================================================
if __name__ == '__main__':
    # === 测试1：死亡案件（冉某江案） ===
    print("=" * 60)
    print("测试1：冉某江等交通肇事案（刑事附带民事·死亡案件）")
    print("=" * 60)

    case1 = {
        'injury_type': 'death',
        'case_type': 'criminal_attached',
        'medical_fee': 102309.83,
        'hospital_days': 11,
        'nursing_days': 11,
        'nursing_custom_daily': 139,
        'lost_work_days': 11,
        'lost_work_daily': 178,
        'traffic_fee': 3000,
        'property_damage': 2000,
        'dependents': [{'age': 75, 'years': 5, 'share_ratio': 0.5}],
        'plaintiff_name': '冉某江',
        'defendant_text': '潘某坤、中国某某财产保险股份有限公司邢台市中心支公司',
        'court': '河北省广宗县人民法院',
    }

    evidence1 = collect_evidence(case1)
    print(f"\n推导出 {len(evidence1)} 项证据：")
    for i, ev in enumerate(evidence1):
        print(f"  {i+1}. {ev['name']} — {ev['purpose']}")

    output1 = generate_evidence_list_docx(
        case1,
        output_path='交通事故案件工作流/证据清单_冉某江案.docx',
        plaintiff_name='冉某江',
        defendant_names=['潘某坤', '中国某某财产保险股份有限公司邢台市中心支公司'],
        court='河北省广宗县人民法院',
    )
    print(f"\n✓ 证据清单已生成: {output1}")

    # 简版文本（用于嵌入起诉状）
    text1 = generate_evidence_list_text(case1)
    print(f"\n简版证据清单（起诉状用）:\n{text1}")

    # === 测试2：伤残案件 ===
    print("\n" + "=" * 60)
    print("测试2：普通伤残案件（纯民事·10级伤残）")
    print("=" * 60)

    case2 = {
        'injury_type': 'disability',
        'case_type': 'civil',
        'medical_fee': 52340.80,
        'hospital_days': 45,
        'nursing_days': 45,
        'nursing_custom_daily': None,
        'lost_work_days': 180,
        'lost_work_daily': 178,
        'traffic_fee': 680,
        'property_damage': 0,
        'dependents': [],
        'plaintiff_name': '李明',
        'defendant_text': '张某某、某某保险股份有限公司',
        'court': '河北省邢台市桥西区人民法院',
    }

    evidence2 = collect_evidence(case2)
    print(f"\n推导出 {len(evidence2)} 项证据：")
    for i, ev in enumerate(evidence2):
        print(f"  {i+1}. {ev['name']} — {ev['purpose']}")

    output2 = generate_evidence_list_docx(
        case2,
        output_path='交通事故案件工作流/证据清单_李明伤残案.docx',
        plaintiff_name='李明',
        defendant_names=['张某某', '某某保险股份有限公司'],
        court='河北省邢台市桥西区人民法院',
    )
    print(f"\n✓ 证据清单已生成: {output2}")
